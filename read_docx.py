# -*- coding: utf-8 -*-
from docx import Document
import os

file_path = r'C:\Users\ips\Desktop\Upan\杭州非遗\定稿-招标文件-杭州市"文化两中心"非遗数字展陈内容制作及相关配套项目.docx'
output_path = r'C:\Users\ips\Desktop\测试AI员工能力\重做投标网站\招标文件内容.txt'

doc = Document(file_path)
text = ''
for para in doc.paragraphs:
    text += para.text + '\n'

with open(output_path, 'w', encoding='utf-8') as f:
    f.write(text)

print(f'共提取 {len(doc.paragraphs)} 段落，{len(text)} 字符')
print(f'已保存到: {output_path}')
print('\n前3000字预览：')
print(text[:3000])
