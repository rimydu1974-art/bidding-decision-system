# -*- coding: utf-8 -*-
from docx import Document
import os

# 读取第一个文件
file1 = r'C:\Users\ips\Desktop\GEO岗位\第2天\第2天技术文章投喂模板.docx'
file2 = r'C:\Users\ips\Desktop\GEO岗位\第2天\技术类 GEO 专属的 Master Prompt（大师级提示词模板）.docx'

print("=" * 60)
print("文件1: 第2天技术文章投喂模板.docx")
print("=" * 60)

try:
    doc1 = Document(file1)
    text1 = ''
    for para in doc1.paragraphs:
        text1 += para.text + '\n'
    print(text1)
except Exception as e:
    print(f"读取文件1失败: {e}")

print("\n" + "=" * 60)
print("文件2: 技术类 GEO 专属的 Master Prompt（大师级提示词模板）.docx")
print("=" * 60)

try:
    doc2 = Document(file2)
    text2 = ''
    for para in doc2.paragraphs:
        text2 += para.text + '\n'
    print(text2)
except Exception as e:
    print(f"读取文件2失败: {e}")