# -*- coding: utf-8 -*-
from docx import Document
import os
import sys

# 设置控制台编码
sys.stdout.reconfigure(encoding='utf-8')

# 使用glob找到的文件路径
file_path = r'C:\Users\ips\Desktop\Upan\杭州非遗\定稿-招标文件-杭州市"文化两中心"非遗数字展陈内容制作及相关配套项目.docx'
output_path = r'C:\Users\ips\Desktop\测试AI员工能力\重做投标网站\招标文件内容.txt'

try:
    doc = Document(file_path)
    text = ''
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + '\n'
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    
    print(f'成功! 共提取 {len(doc.paragraphs)} 段落，{len(text)} 字符')
    print(f'已保存到: {output_path}')
    print(f'\n前3000字预览：')
    print(text[:3000])
except Exception as e:
    print(f'错误: {e}')
    # 尝试列出目录
    dir_path = r'C:\Users\ips\Desktop\Upan\杭州非遗'
    if os.path.exists(dir_path):
        print(f'\n目录内容:')
        for f in os.listdir(dir_path):
            if '招标' in f:
                print(f'  >>> {f}')
